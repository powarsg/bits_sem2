"""Build designPS04_G116.docx (editable), contribution xlsx, and submission zip."""

import os
import zipfile
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

BASE = os.path.dirname(os.path.abspath(__file__))
IMG = os.path.join(BASE, "images")
OUT_DOCX = os.path.join(BASE, "designPS04_G116.docx")
OUT_XLSX = os.path.join(BASE, "G116_Contribution.xlsx")


def _style_table(table, header=True):
    if header and table.rows:
        for cell in table.rows[0].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = str(h)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    _style_table(table)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def add_image(doc, filename, width=3.0):
    path = os.path.join(IMG, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.add_paragraph()
    else:
        add_para(doc, f"[Image not found: {filename}]")


def build_docx():
    doc = Document()

    # Title block
    title = doc.add_heading("Defence Drone Navigation using GBFS and A* Search", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Course: AIMLCZG557/AECLZG557  |  Assignment: PS4  |  Group: G116")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(80, 80, 80)
    doc.add_paragraph()

    # 1. Project Summary
    add_heading(doc, "1. Project Summary", 2)
    add_para(
        doc,
        "An autonomous UAV navigates an 8×8 grid from Start (0,0) to Goal (6,7), avoiding "
        "No-Fly Zones (N) and minimizing Weather Hazard (W) penalties. Implemented: GBFS with "
        "h1 (Euclidean), GBFS with h2 (Bounding-Box Risk-Weighted), and A* with h1 for comparison.",
    )

    # 2. Environment
    add_heading(doc, "2. Environment & Problem Modelling", 2)
    add_table(
        doc,
        ["Symbol", "Description", "Cost"],
        [
            ["S", "Start (0,0)", "1"],
            ["E", "Goal (6,7)", "1"],
            [".", "Passable", "1"],
            ["W", "Weather Hazard", "4"],
            ["N", "No-Fly Zone", "Blocked"],
        ],
    )
    add_para(doc, "Grid: 8×8 | No-Fly: 8 | Weather: 7 | Movement: N,E,S,W (tie-break order)")

    # 3. PEAS
    add_heading(doc, "3. PEAS Components", 2)
    add_table(
        doc,
        ["Component", "Description"],
        [
            ["Performance", "Minimize path cost; reach goal; avoid N zones; reduce W penalties"],
            ["Environment", "8×8 grid, static, fully observable, deterministic, single agent"],
            ["Actuators", "Orthogonal movement (N/S/E/W); navigation controller"],
            ["Sensors", "GPS position; grid map; obstacle/weather/goal detection"],
        ],
    )

    # 4. Alternate modelling
    add_heading(doc, "4. Alternate Modelling Approach", 2)
    add_para(
        doc,
        "Primary model: 2D matrix grid with O(1) cell lookup. Alternate: explicit graph G=(V,E) "
        "where V = passable cells (56 nodes) and E = orthogonal edges with transition costs. "
        "Performance: matrix is faster for neighbour expansion on dense grids; graph model uses "
        "less memory on sparse/large maps and scales better when the grid is not rectangular. "
        "For this 8×8 assignment both are equivalent; matrix was chosen for simplicity.",
    )

    # 5. Architecture
    add_heading(doc, "5. System Architecture", 2)
    add_table(
        doc,
        ["Component", "Choice", "Reason"],
        [
            ["Environment", "2D Matrix", "O(1) access"],
            ["OPEN", "Priority Queue (heap)", "Efficient min-h extraction"],
            ["CLOSED", "Hash Set", "O(1) revisit check"],
            ["Path", "Parent dict", "O(path) reconstruction"],
        ],
    )

    # 6. Heuristics
    add_heading(doc, "6. Heuristic Engineering & Execution Flow", 2)
    add_para(doc, "Heuristic h1 – Euclidean Distance", bold=True)
    add_para(doc, "h1(n) = √((xr−xg)² + (yr−yg)²). Fast, goal-directed, admissible for A*.")
    add_para(doc, "Heuristic h2 – Bounding-Box Risk-Weighted", bold=True)
    add_para(
        doc,
        "h2(n) = Manhattan(n,goal) × avg(cell costs in bounding box from n to goal). "
        "Future-aware: penalises paths through high-cost hazard regions. Example: h2(3,3)=15.40.",
    )
    add_para(doc, "GBFS Algorithm", bold=True)
    add_para(doc, "1. Pop min-h node from OPEN  2. If goal, stop  3. Add to CLOSED  4. Expand neighbours  5. Repeat")

    add_para(doc, "Execution Flow (Sequence Diagram)", bold=True)
    for line in [
        "1. Read inputPS04.txt  →  Load 8×8 grid, start (0,0), goal (6,7)",
        "2. Run GBFS-h1         →  Euclidean heuristic search, record path & metrics",
        "3. Run GBFS-h2         →  Bounding-box heuristic search, record path & metrics",
        "4. Run A*-h1           →  Optimal path benchmark with f(n)=g(n)+h(n)",
        "5. Write outputPS04.txt → Metrics table, PEAS, trap check, complexity",
        "6. Generate PNG charts  → Grid paths, comparison, heuristic plots",
    ]:
        add_para(doc, line)

    add_para(doc, "Data Structures & Tie-Breaking", bold=True)
    add_table(
        doc,
        ["Structure", "Role", "Complexity"],
        [
            ["heapq (OPEN)", "Frontier priority queue", "push/pop O(log n)"],
            ["set (CLOSED)", "Explored nodes", "add/lookup O(1)"],
            ["dict (parent)", "Path reconstruction", "O(1) lookup"],
        ],
    )
    add_para(doc, "Node priority on equal h: North > East > South > West (counter tie-break in heap).")

    # 7. Results
    add_heading(doc, "7. Results & Comparison (outputPS04.txt)", 2)
    add_para(doc, "Metrics Table", bold=True)
    add_table(
        doc,
        ["Metric", "Description", "GBFS-h1", "GBFS-h2", "A*-h1"],
        [
            ["Nodes Expanded", "Total explored nodes", "14", "14", "29"],
            ["Runtime", "Execution time (ms)", "0.1390", "0.1910", "0.2279"],
            ["Memory Usage", "OPEN + CLOSED size", "23", "28", "40"],
            ["Total Path Cost", "Sum of transition costs", "17", "17", "14"],
            ["Path Length", "Total moves", "13", "13", "13"],
            ["Heuristic", "Heuristic value at start", "9.2195", "29.4821", "9.2195"],
            ["Weather Crossed", "Weather zones entered", "1", "1", "0"],
            ["Weather Penalty", "Total weather cost", "4", "4", "0"],
        ],
    )

    add_para(doc, "GBFS vs A* Comparison (Heuristic h1)", bold=True)
    add_table(
        doc,
        ["Criterion", "GBFS (h1)", "A* (h1)"],
        [
            ["Complete?", "Yes (with closed set)", "Yes"],
            ["Optimal?", "No", "Yes"],
            ["Nodes Expanded", "14", "29"],
            ["Runtime (ms)", "0.1390", "0.2279"],
            ["Memory Usage", "23", "40"],
            ["Total Path Cost", "17", "14"],
            ["Path Length", "13", "13"],
            ["Faster to goal?", "Yes (fewer nodes, lower ms)", "No (more exploration)"],
        ],
    )

    add_para(doc, "Trap Identification (Expectation 6)", bold=True)
    add_table(
        doc,
        ["Heuristic", "H-Value", "Trapped at Node", "Iteration", "Escaped To", "Escape Iter."],
        [["—", "—", "No traps detected", "—", "—", "—"]],
    )
    add_para(
        doc,
        "Both GBFS heuristics found 13-move paths costing 17 (1 weather zone). "
        "A* found optimal cost 14 with zero weather penalties.",
    )

    # 8. Analysis & charts
    add_heading(doc, "8. Analysis, Visualizations & Conclusion", 2)
    add_para(doc, "Heuristic & Algorithm Comparison Charts", bold=True)
    add_image(doc, "comparison_chart.png", 5.5)
    add_image(doc, "heuristic_progression.png", 5.5)
    add_image(doc, "heuristic_vs_time.png", 5.5)
    add_image(doc, "grid_gbfs_h1.png", 3.0)
    add_image(doc, "grid_astar_h1.png", 3.0)

    add_para(doc, "Complexity Analysis", bold=True)
    add_table(
        doc,
        ["Algorithm", "Time", "Space", "Nodes", "Notes"],
        [
            ["GBFS h1", "O(V log V)", "O(V)", "14", "Fastest; not optimal"],
            ["GBFS h2", "O(V log V)", "O(V)", "14", "Risk-aware; higher h2 compute"],
            ["A* h1", "O(V log V)", "O(V)", "29", "Optimal path; more exploration"],
        ],
    )

    add_para(doc, "Heuristic Analysis", bold=True)
    add_para(
        doc,
        "h1 (Euclidean): fastest runtime (0.139 ms), simple geometry, ignores hazard layout. "
        "h2 (Bounding-Box): higher start h=29.48, accounts for future hazard costs; same path cost "
        "on this grid. h1 is better for speed; h2 is better for hazard-aware routing on complex maps.",
    )

    add_heading(doc, "9. Conclusion", 2)
    add_para(
        doc,
        "GBFS is complete with a closed set on finite grids but not optimal — it found cost-17 paths "
        "vs A* optimal cost-14. GBFS-h1 is fastest (14 nodes, 0.139 ms); A* guarantees optimality at "
        "higher cost (29 nodes, 0.228 ms). GBFS-h2 balances efficiency with risk awareness. "
        "Recommendation: A* for mission-critical optimality; GBFS-h1 for real-time response; "
        "GBFS-h2 for hazard-dense environments.",
    )
    add_para(
        doc,
        "Additional visualizations in output: grid_initial, grid_gbfs_h2, heuristic_heatmaps, nodes_explored.",
    )

    doc.save(OUT_DOCX)
    print(f"Created {OUT_DOCX}")
    return OUT_DOCX


def build_contribution_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "Contribution"
    headers = ["Student Registration Number", "Name", "Percentage of contribution out of 100%"]
    header_fill = PatternFill("solid", fgColor="4472C4")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    thin = Side(style="thin")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", wrap_text=True)
        cell.border = border

    placeholders = [
        ("<Reg No 1>", "<Member Name 1>", 50),
        ("<Reg No 2>", "<Member Name 2>", 50),
    ]
    for row_idx, (reg, name, pct) in enumerate(placeholders, 2):
        ws.cell(row=row_idx, column=1, value=reg).border = border
        ws.cell(row=row_idx, column=2, value=name).border = border
        c = ws.cell(row=row_idx, column=3, value=pct)
        c.border = border
        c.alignment = Alignment(horizontal="center")

    ws.column_dimensions["A"].width = 32
    ws.column_dimensions["B"].width = 28
    ws.column_dimensions["C"].width = 38
    ws.row_dimensions[1].height = 30
    wb.save(OUT_XLSX)
    print(f"Created {OUT_XLSX} (update member details before submitting)")


def build_zip():
    ts = datetime.now().strftime("%Y%m%d%H%M")
    zip_name = os.path.join(BASE, f"G116_A1_PS04_{ts}.zip")
    files = [
        "designPS04_G116.docx",
        "G116_Contribution.xlsx",
        "inputPS04.txt",
        "outputPS04.txt",
        "team_116.py",
    ]
    with zipfile.ZipFile(zip_name, "w", zipfile.ZIP_DEFLATED) as zf:
        for fname in files:
            path = os.path.join(BASE, fname)
            if os.path.exists(path):
                zf.write(path, fname)
            else:
                print(f"WARNING: missing {fname}")
    print(f"Created {zip_name}")
    return zip_name


if __name__ == "__main__":
    build_docx()
    build_contribution_xlsx()
    build_zip()
